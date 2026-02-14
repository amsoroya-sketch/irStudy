#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Set default font
style = doc.styles["Normal"]
font = style.font
font.name = "Arial"
font.size = Pt(11)

# Title
title = doc.add_heading("Hepatitis B Vaccination Completion Plan", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.font.size = Pt(18)
title_run.font.bold = True

# Subtitle
subtitle = doc.add_heading("For NSW Health – IMG Clinical Readiness Program (ICRP)", level=2)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.runs[0]
subtitle_run.font.size = Pt(14)
subtitle_run.font.bold = True

# Add some space
doc.add_paragraph()

# Employee Name
p = doc.add_paragraph()
p.add_run("Employee Name: ").bold = True
p.add_run("Dr Iram Asim")

# Date of Birth
p = doc.add_paragraph()
p.add_run("Date of Birth: ").bold = True
p.add_run("_" * 40)

# Add space
doc.add_paragraph()

# Purpose
p = doc.add_paragraph()
p.add_run("Purpose: ").bold = True
p.add_run("NSW Health Employment – ICRP Checking Stage")

# Add space
doc.add_paragraph()

# Confirmation text
doc.add_paragraph(
    "This is to confirm that the above-named individual is currently undertaking a Hepatitis B vaccination course in accordance with NSW Health Occupational Assessment, Screening and Vaccination (OASV) requirements."
)

# Add horizontal line (using a border on an empty paragraph)
p = doc.add_paragraph()
p_format = p.paragraph_format
p_format.space_before = Pt(12)
p_format.space_after = Pt(12)

# Vaccination Schedule Header
heading = doc.add_heading("Hepatitis B Vaccination Schedule", level=3)
heading_run = heading.runs[0]
heading_run.font.size = Pt(12)
heading_run.font.bold = True

# Dose 1
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
p.add_run("Dose 1: ").bold = True
p.add_run("✓ ").bold = True
p.add_run("Administered on ")
p.add_run("16 / 12 / 2025").bold = True

# Dose 2
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
p.add_run("Dose 2: ").bold = True
p.add_run("☐ ")
p.add_run("Planned on ")
p.add_run("16 / 01 / 2026").bold = True

# Dose 3
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
p.add_run("Dose 3: ").bold = True
p.add_run("☐ ")
p.add_run("Planned on ")
p.add_run("16 / 06 / 2026").bold = True

# Add space
doc.add_paragraph()

# Post-vaccination serology
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
p.add_run("Post-vaccination serology (Anti-HBs):").bold = True

p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(1.0)
p.add_run("☐ ")
p.add_run("Planned on ")
p.add_run("21 / 07 / 2026").bold = True

p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(1.0)
run = p.add_run("(4–8 weeks after Dose 3)")
run.font.size = Pt(10)
run.italic = True

# Add space
doc.add_paragraph()

# Confirmation statement
doc.add_paragraph(
    "All remaining doses and post-vaccination serology will be completed within recommended timeframes, and documentary evidence will be provided to NSW Health."
)

# Add horizontal line
p = doc.add_paragraph()
p_format = p.paragraph_format
p_format.space_before = Pt(12)
p_format.space_after = Pt(12)

# Signature section
doc.add_paragraph()

# GP Name
p = doc.add_paragraph()
p.add_run("GP Name: ").bold = True
p.add_run("_" * 50)

# Provider Number
p = doc.add_paragraph()
p.add_run("Provider Number: ").bold = True
p.add_run("_" * 40)

# Clinic Name & Stamp
p = doc.add_paragraph()
p.add_run("Clinic Name & Stamp: ").bold = True
p.add_run("_" * 50)

# Signature
p = doc.add_paragraph()
p.add_run("Signature: ").bold = True
p.add_run("_" * 50)

# Date
p = doc.add_paragraph()
p.add_run("Date: ").bold = True
p.add_run("___ / ___ / _____")

# Save the document
doc.save("/home/dev/Development/irStudy/Hepatitis_B_Vaccination_Plan.docx")
print("Word document created successfully: Hepatitis_B_Vaccination_Plan.docx")
